from docx import Document

class WordProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)
        
        # Create mappings for quick lookup
        self._create_element_mappings()
        
    def _create_element_mappings(self):
        """Create mappings between elements and their objects for quick lookup."""
        # Map paragraph elements to paragraph objects
        self.paragraph_map = {
            para._element: para 
            for para in self.document.paragraphs
        }
        
        # Map table elements to table objects
        self.table_map = {
            table._element: table 
            for table in self.document.tables
        }

    def _format_table(self, table):
        """Convert table to string format."""
        table_string = ''
        for row in table.rows:
            # Convert each cell to text, handle empty cells
            cells = [cell.text.strip() if cell.text.strip() else 'None' for cell in row.cells]
            table_string += '|' + '|'.join(cells) + '|\n'
        return table_string

    def process_file(self):
        """Process document maintaining the order of paragraphs and tables."""
        content = []
        
        # Get all elements in the document
        elements = self.document.element.body

        # iterate through each element to check its tag
        for element in elements:
            if element.tag.endswith('p'):  # Paragraph
                # Direct lookup instead of loop
                if element in self.paragraph_map:
                    para = self.paragraph_map[element]
                    if para.text.strip():
                        content.append(para.text)
                        
            elif element.tag.endswith('tbl'):  # Table
                # Direct lookup instead of loop
                if element in self.table_map:
                    table = self.table_map[element]
                    content.append("\n" + self._format_table(table) + "\n")
        
        # Join all content
        return '\n'.join(content)
