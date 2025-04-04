import logging
import os
import re
import shutil
import tempfile
import uuid
from io import BytesIO
from typing import List

import chardet
import docx
import fitz  # PyMuPDF
import pandas as pd
import requests
from config import get_supabase_client
from dotenv import load_dotenv
from langchain.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from langchain.text_splitter import (
    MarkdownTextSplitter,
    RecursiveCharacterTextSplitter,
    TokenTextSplitter,
)
from langchain.vectorstores import FAISS
from models.request.rag_request import ProcessDocumentsRequest
from models.response.rag_response import ProcessDocumentsResponse
from models.response.response_wrapper import ErrorResponse, SuccessResponse
from PyPDF2 import PdfReader
from services.facade.rag_service import RAGService
from supabase import Client

BUCKET_NAME = "DOCUMENTS"

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

class RAGServiceImpl(RAGService):
    def __init__(self):
        load_dotenv()
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("Please set OPENAI_API_KEY in the .env file")
        self.embeddings = OpenAIEmbeddings()
        # Enhanced text splitters
        self.text_splitters = {
            "default": RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len
            ),
            "markdown": MarkdownTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            ),
            "token": TokenTextSplitter(
                chunk_size=500,
                chunk_overlap=50
            )
        }
        
        self.text_splitter = self.text_splitters["default"]
        self.vector_store = None


    def _determine_text_splitter(self, file_type: str) -> RecursiveCharacterTextSplitter:
        """Determine the best text splitter based on file type"""
        if file_type in ['.md', '.markdown']:
            return self.text_splitters["markdown"]
        elif file_type in ['.pdf', '.docx', '.doc']:
            return self.text_splitters["token"]
        else:
            return self.text_splitters["default"]


    def _save_vector_store(self, user_id: str, chatbot_id: str, supabase: Client):
        """Save vector store to Supabase storage"""
        if not self.vector_store:
            logging.warning("No vector store to save")
            return
        try:
            # Create a temporary directory to save files
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save vector store locally first
                self.vector_store.save_local(temp_dir)
                
                faiss_path = os.path.join(temp_dir, "index.faiss")
                pkl_path = os.path.join(temp_dir, "index.pkl")
                
                # Define storage paths
                storage_faiss_path = f"{user_id}/{chatbot_id}/rag-vector/index.faiss"
                storage_pkl_path = f"{user_id}/{chatbot_id}/rag-vector/index.pkl"
                
                # Check if the files already exist in the bucket
                result = supabase.table('documents') \
                    .select('id') \
                    .eq('chatbot_id', chatbot_id) \
                    .eq('file_name', 'index.faiss') \
                    .execute()
                
                if result.data:
                    logging.info(f"Vector store files already exist in the bucket, removing them first")
                    supabase.storage.from_(BUCKET_NAME).remove([storage_faiss_path, storage_pkl_path])
                    supabase.table('documents').delete().eq('chatbot_id', chatbot_id).eq('file_name', 'index.faiss').execute()
                    supabase.table('documents').delete().eq('chatbot_id', chatbot_id).eq('file_name', 'index.pkl').execute()

                # Upload files directly to bucket
                with open(faiss_path, 'rb') as f:
                    supabase.storage.from_(BUCKET_NAME).upload(storage_faiss_path, f)
                with open(pkl_path, 'rb') as f:
                    supabase.storage.from_(BUCKET_NAME).upload(storage_pkl_path, f)
                logging.info(f"Vector store saved to Supabase bucket for chatbot {chatbot_id}")

                # add two rows to the documents table
                supabase.table('documents').insert({
                    'id': str(uuid.uuid4()),
                    'chatbot_id': chatbot_id,
                    'file_name': 'index.faiss',
                    'file_type': 'faiss',
                    'is_processed': True,
                    'bucket_path': storage_faiss_path
                }).execute()

                supabase.table('documents').insert({
                    'id': str(uuid.uuid4()),
                    'chatbot_id': chatbot_id,
                    'file_name': 'index.pkl',
                    'file_type': 'pkl',
                    'is_processed': True,
                    'bucket_path': storage_pkl_path
                }).execute()
                
                logging.info(f"Added two rows to the documents table for chatbot {chatbot_id}")
                
        except Exception as e:
            logging.error(f"Error saving vector store to Supabase: {str(e)}")
            raise


    def _process_url_document(self, url: str) -> List[Document]:
        """Process document from URL"""
        try:
            response = requests.get(url)
            response.raise_for_status()
            
            filename = url.split('/')[-1].split('?')[0]
            file_content = BytesIO(response.content)
            file_type = os.path.splitext(filename)[1].lower()
            self.text_splitter = self._determine_text_splitter(file_type)

            # Create a temp directory for image extraction if needed
            temp_dir = None
            if file_type in ['.pdf', '.docx', '.doc']:
                temp_dir = tempfile.mkdtemp()
                
            try: 
                if filename.endswith('.pdf'):
                    try:
                        # For text extraction
                        pdf_reader = PdfReader(file_content)
                        docs = []
                        
                        # Try to use PyMuPDF for better extraction if available
                        try:
                            # Save to a temporary file for PyMuPDF (it works better with files)
                            temp_file_path = os.path.join(temp_dir, filename) if temp_dir else None
                            if temp_file_path:
                                file_content.seek(0)
                                with open(temp_file_path, 'wb') as f:
                                    f.write(file_content.read())
                                
                                # Use PyMuPDF for advanced extraction
                                pdf_document = fitz.open(temp_file_path)
                                
                                # Process each page
                                for page_idx, page in enumerate(pdf_document):
                                    # Extract text with better formatting
                                    page_text = page.get_text("text")
                                    
                                    # Create a document for each page
                                    if page_text.strip():
                                        docs.append(Document(
                                            page_content=page_text,
                                            metadata={
                                                "source": url,
                                                "page": page_idx + 1,
                                                "total_pages": len(pdf_document),
                                                "file_type": "pdf",
                                                "file_name": filename
                                            }
                                        ))
                                
                                pdf_document.close()
                            
                        except ImportError:
                            # Fall back to PyPDF2 if PyMuPDF is not available
                            for idx, page in enumerate(pdf_reader.pages):
                                text = page.extract_text()
                                if text.strip():  # Skip empty pages
                                    docs.append(Document(
                                        page_content=text,
                                        metadata={
                                            "source": url,
                                            "page": idx + 1,
                                            "total_pages": len(pdf_reader.pages),
                                            "file_type": "pdf",
                                            "file_name": filename
                                        }
                                    ))
                        
                        logging.info(f"Successfully processed PDF file from URL: {url}")
                        
                    except Exception as pdf_error:
                        logging.error(f"Error in PDF processing: {str(pdf_error)}")
                        # Fall back to basic PDF extraction
                        for idx, page in enumerate(pdf_reader.pages):
                            text = page.extract_text()
                            if text.strip():  # Skip empty pages
                                docs.append(Document(
                                    page_content=text,
                                    metadata={
                                        "source": url,
                                        "page": idx + 1,
                                        "total_pages": len(pdf_reader.pages),
                                        "file_type": "pdf",
                                        "file_name": filename
                                    }
                                ))
                    
                elif filename.endswith('.txt'):
                    content = file_content.read().decode('utf-8')
                    processed_content = content.replace('\r\n', '\n')  # Normalize line endings
                    # Create paragraphs from double line breaks
                    paragraphs = [p.strip() for p in processed_content.split('\n\n') if p.strip()]
                    formatted_content = '\n\n'.join(paragraphs)
                    docs = [Document(
                            page_content=formatted_content,
                            metadata={
                                "source": url,
                                "file_type": file_type,
                                "file_name": filename,
                                "paragraph_count": len(paragraphs)
                            }
                        )]
                    logging.info(f"Successfully processed text file from URL: {url}")
                    
                elif filename.endswith('.csv'):
                    try:
                        # First try to infer the encoding
                        file_content.seek(0)
                        sample = file_content.read(min(1024 * 10, len(file_content.getbuffer())))
                        
                        # Chardet for encoding detection
                        try:
                            detected = chardet.detect(sample)
                            encoding = detected.get('encoding', 'utf-8')
                        except ImportError:
                            encoding = 'utf-8'
                        
                        # Reset file position
                        file_content.seek(0)
                        
                        # Try to detect delimiter
                        sample_str = sample.decode(encoding, errors='replace')
                        possible_delimiters = [',', ';', '\t', '|']
                        delimiter_counts = {delim: sample_str.count(delim) for delim in possible_delimiters}
                        delimiter = max(delimiter_counts, key=delimiter_counts.get)
                        
                        # Read CSV with pandas
                        df = pd.read_csv(file_content, delimiter=delimiter, encoding=encoding, on_bad_lines='skip')
                        
                        docs = []
                        
                        # Process by chunks to preserve related rows
                        chunk_size = 20  # Adjust based on your needs
                        for i in range(0, len(df), chunk_size):
                            chunk = df.iloc[i:i+chunk_size]
                            
                            # Format as markdown table for better structure
                            md_table = "| " + " | ".join(str(col) for col in df.columns) + " |\n"
                            md_table += "| " + " | ".join("---" for _ in df.columns) + " |\n"
                            
                            for _, row in chunk.iterrows():
                                md_table += "| " + " | ".join(str(val) for val in row) + " |\n"
                            
                            # Also include row-by-row data for easier searching
                            row_data = []
                            for idx, row in chunk.iterrows():
                                row_str = f"Row #{idx + 1}:\n"
                                for col in df.columns:
                                    row_str += f"  {col}: {row[col]}\n"
                                row_data.append(row_str)
                            
                            # Combine table and row data
                            content = md_table + "\n\n" + "\n\n".join(row_data)
                            
                            doc = Document(
                                page_content=content,
                                metadata={
                                    "source": url,
                                    "chunk_start": i,
                                    "chunk_end": min(i+chunk_size, len(df)),
                                    "total_rows": len(df),
                                    "columns": list(df.columns),
                                    "file_type": "csv",
                                    "file_name": filename
                                }
                            )
                            docs.append(doc)
                                
                    except Exception as csv_error:
                        logging.warning(f"Error in advanced CSV processing: {str(csv_error)}. Falling back to basic processing.")
                        # Fall back to basic CSV processing
                        file_content.seek(0)
                        df = pd.read_csv(file_content)
                        docs = []
                        for idx, row in df.iterrows():
                            content = f"Record #{idx + 1}\n"
                            for col in df.columns:
                                content += f"{col}: {row[col]}\n"
                            doc = Document(
                                page_content=content,
                                metadata={
                                    "source": url, 
                                    "row": idx,
                                    "file_type": "csv",
                                    "file_name": filename
                                }
                            )
                            docs.append(doc)
                    
                    logging.info(f"Successfully processed CSV file from URL: {url}")
                    
                elif filename.endswith('.md'):
                    content = file_content.read().decode('utf-8')
                        
                    # Process headers and structure
                    sections = []
                    current_section = {"title": "", "content": "", "level": 0}
                    lines = content.split('\n')
                    
                    for line in lines:
                        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
                        if header_match:
                            # If we were building a section, save it
                            if current_section["content"]:
                                sections.append(current_section)
                            
                            # Start a new section
                            level = len(header_match.group(1))
                            title = header_match.group(2)
                            current_section = {
                                "title": title,
                                "content": line + "\n",
                                "level": level
                            }
                        else:
                            # Add to current section
                            current_section["content"] += line + "\n"
                    
                    # Add the last section
                    if current_section["content"]:
                        sections.append(current_section)
                    
                    # Create documents based on sections
                    if sections:
                        docs = []
                        for i, section in enumerate(sections):
                            # Create a document for each major section
                            docs.append(Document(
                                page_content=section["content"],
                                metadata={
                                    "source": url,
                                    "file_type": "markdown",
                                    "file_name": filename,
                                    "section_title": section["title"],
                                    "section_level": section["level"],
                                    "section_index": i,
                                    "total_sections": len(sections)
                                }
                            ))
                    else:
                        # No sections found, treat as one document
                        docs = [Document(
                            page_content=content,
                            metadata={
                                "source": url,
                                "file_type": "markdown",
                                "file_name": filename
                            }
                        )]
                    
                    logging.info(f"Successfully processed markdown file from URL: {url}")
                    
                elif filename.endswith(('.xlsx', '.xls')):
                    try:
                        # Read all sheets
                        file_content.seek(0)
                        excel_file = pd.ExcelFile(file_content)
                        sheet_names = excel_file.sheet_names
                        
                        docs = []
                        
                        for sheet_name in sheet_names:
                            df = pd.read_excel(excel_file, sheet_name=sheet_name)
                            
                            # Skip empty sheets
                            if df.empty:
                                continue
                                
                            # Process by chunks to preserve context
                            chunk_size = 20  # Adjust based on your needs
                            for i in range(0, len(df), chunk_size):
                                chunk = df.iloc[i:i+chunk_size]
                                
                                # Format as markdown table for better structure
                                md_table = f"## Sheet: {sheet_name}\n\n"
                                md_table += "| " + " | ".join(str(col) for col in df.columns) + " |\n"
                                md_table += "| " + " | ".join("---" for _ in df.columns) + " |\n"
                                
                                for _, row in chunk.iterrows():
                                    md_table += "| " + " | ".join(str(val) for val in row) + " |\n"
                                
                                # Also include row-by-row data for easier searching
                                row_data = []
                                for idx, row in chunk.iterrows():
                                    row_str = f"Row #{idx + 1}:\n"
                                    for col in df.columns:
                                        row_str += f"  {col}: {row[col]}\n"
                                    row_data.append(row_str)
                                
                                # Combine table and row data
                                content = md_table + "\n\n" + "\n\n".join(row_data)
                                
                                doc = Document(
                                    page_content=content,
                                    metadata={
                                        "source": url,
                                        "sheet_name": sheet_name,
                                        "chunk_start": i,
                                        "chunk_end": min(i+chunk_size, len(df)),
                                        "total_sheets": len(sheet_names),
                                        "total_rows": len(df),
                                        "columns": list(df.columns),
                                        "file_type": "excel",
                                        "file_name": filename
                                    }
                                )
                                docs.append(doc)
                        
                    except Exception as excel_error:
                        logging.warning(f"Error in advanced Excel processing: {str(excel_error)}. Falling back to basic processing.")
                        # Fall back to basic Excel processing
                        file_content.seek(0)
                        df = pd.read_excel(file_content)
                        docs = []
                        for idx, row in df.iterrows():
                            content = f"Record #{idx + 1}\n"
                            for col in df.columns:
                                content += f"{col}: {row[col]}\n"
                            doc = Document(
                                page_content=content,
                                metadata={
                                    "source": url, 
                                    "row": idx,
                                    "file_type": "excel",
                                    "file_name": filename
                                }
                            )
                            docs.append(doc)
                    
                    logging.info(f"Successfully processed Excel file from URL: {url}")
                    
                elif filename.endswith(('.docx', '.doc')):
                    try:
                        # Save to a temporary file for better processing
                        temp_file_path = os.path.join(temp_dir, filename) if temp_dir else None
                        if temp_file_path:
                            file_content.seek(0)
                            with open(temp_file_path, 'wb') as f:
                                f.write(file_content.read())
                            
                            # Process document with python-docx
                            doc = docx.Document(temp_file_path)
                            
                            # Extract document structure
                            content = []
                            headings = []
                            
                            # Process paragraphs and extract structure
                            for para in doc.paragraphs:
                                if para.text.strip():
                                    # Check if it's a heading
                                    if para.style.name.startswith('Heading'):
                                        heading_level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
                                        heading = '#' * heading_level + ' ' + para.text
                                        content.append(heading)
                                        headings.append({
                                            "text": para.text,
                                            "level": heading_level
                                        })
                                    else:
                                        # Regular paragraph
                                        content.append(para.text)
                            
                            # Process tables
                            for table in doc.tables:
                                table_data = []
                                table_markdown = "| "
                                
                                # Extract header row if exists
                                header_row = []
                                if table.rows:
                                    for cell in table.rows[0].cells:
                                        header_row.append(cell.text.strip())
                                    table_markdown += " | ".join(header_row) + " |\n| "
                                    table_markdown += " | ".join(["---"] * len(header_row)) + " |\n"
                                
                                # Extract data rows
                                for i, row in enumerate(table.rows):
                                    if i == 0 and header_row:  # Skip header in data extraction
                                        continue
                                        
                                    row_data = []
                                    for cell in row.cells:
                                        row_data.append(cell.text.strip())
                                    
                                    table_data.append(row_data)
                                    table_markdown += "| " + " | ".join(row_data) + " |\n"
                                
                                # Add table to content
                                content.append(table_markdown)
                            
                            # Combine all content with proper spacing
                            full_content = '\n\n'.join(content)
                            
                            # Create document with enhanced metadata
                            docs = [Document(
                                page_content=full_content,
                                metadata={
                                    "source": url,
                                    "file_type": "word",
                                    "file_name": filename,
                                    "headings": headings,
                                    "contains_tables": len(doc.tables) > 0,
                                    "table_count": len(doc.tables)
                                }
                            )]
                            
                        else:
                            # Fall back if temp file couldn't be created
                            file_content.seek(0)
                            doc = docx.Document(file_content)
                            content = []
                            
                            # Extract text from paragraphs
                            for para in doc.paragraphs:
                                if para.text.strip():
                                    content.append(para.text)
                            
                            # Process tables
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = [cell.text.strip() for cell in row.cells]
                                    if any(row_text):
                                        content.append(" | ".join(row_text))
                            
                            # Join content
                            full_content = '\n\n'.join(content)
                            
                            # Create basic document
                            docs = [Document(
                                page_content=full_content,
                                metadata={
                                    "source": url,
                                    "file_type": "word",
                                    "file_name": filename
                                }
                            )]
                    
                    except Exception as docx_error:
                        logging.warning(f"Error in advanced DOCX processing: {str(docx_error)}. Falling back to basic processing.")
                        # Fall back to basic Word processing
                        file_content.seek(0)
                        doc = docx.Document(file_content)
                        content = []
                        
                        # Extract text from paragraphs
                        for para in doc.paragraphs:
                            if para.text.strip():
                                content.append(para.text)
                        
                        # Process tables
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = [cell.text.strip() for cell in row.cells]
                                if any(row_text):
                                    content.append(" | ".join(row_text))
                        
                        # Join content
                        full_content = '\n'.join(content)
                        
                        # Create basic document
                        docs = [Document(
                            page_content=full_content,
                            metadata={
                                "source": url,
                                "file_type": "word",
                                "file_name": filename
                            }
                        )]
                    
                    logging.info(f"Successfully processed Word file from URL: {url}")
                    
                else:
                    logging.warning(f"Unsupported file type for URL: {url}")
                    return []
                
                return docs
        
            finally:
                # Clean up temp directory if used
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)

        except Exception as e:
            logging.error(f"Error processing file from URL {url}: {str(e)}")
            return []


    def process_documents_from_urls(self, user_id: str, user_token: str, data: ProcessDocumentsRequest) -> tuple[dict, int]:
        """Process documents from URLs (from the documents table in Supabase)"""
        supabase = get_supabase_client(user_token)
        chatbot_id = data.chatbot_id
        # Get document URLs from documents table for newly uploaded documents (is_processed = False)
        try:
            result = supabase.table('documents') \
                .select('bucket_path') \
                .eq('chatbot_id', chatbot_id) \
                .eq('is_processed', False) \
                .execute()
                
            if not result.data:
                return ErrorResponse(
                    message="No unprocessed documents found for this chatbot"
                ).model_dump(), 404
            
            documents = []
            failed_urls = []

            for doc in result.data:
                bucket_path = doc['bucket_path']
                # Get url from bucket_path
                url = supabase.storage.from_(BUCKET_NAME).create_signed_url(bucket_path, 3600)['signedURL']
                # Process the file url
                docs = self._process_url_document(url)
                if docs:
                    documents.extend(docs)
                else:
                    failed_urls.append(str(url))

            if not documents:
                return ErrorResponse(
                    message="No documents were successfully processed",
                    data={"failed_urls": failed_urls}
                ).model_dump(), 400

            logging.info("Starting document splitting...")
            splits = self.text_splitter.split_documents(documents)
            logging.info(f"Document splitting completed, generated {len(splits)} segments")

            # If there is an existing vector store, append new documents to it
            if self.vector_store:
                self.vector_store.add_documents(splits)
                logging.info("Appended new documents to the existing vector store")
            # If there is no existing vector store, create a new one
            else:
                self.vector_store = FAISS.from_documents(documents=splits, embedding=self.embeddings)
                logging.info("Created a new vector store with the new documents")
            
            # Update is_processed to True for the newly processed documents
            supabase.table('documents') \
                .update({'is_processed': True}) \
                .eq('chatbot_id', chatbot_id) \
                .execute()
            
            # Save the vector store to Supabase bucket
            self._save_vector_store(user_id, chatbot_id, supabase)
            
            return SuccessResponse(
                data=ProcessDocumentsResponse(
                    processed_count=len(documents),
                    failed_urls=failed_urls
                ).model_dump(),
                message="Documents processed successfully"
            ).model_dump(), 200
            
        except Exception as e:
            logging.error(f"Error processing documents from URLs: {str(e)}")
            return ErrorResponse(
                message=str(e)
            ).model_dump(), 500